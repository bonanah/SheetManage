from docx import Document
from manage_db import get_problems_by_filter

def generate_test(chapter=None, difficulty=None, output_file="output/new_test.docx"):
    # 워드 문서 생성
    doc = Document()
    doc.add_heading("시험지", level=1)

    # 문제 조회
    problems = get_problems_by_filter(chapter=chapter, difficulty=difficulty)
    if not problems:
        print("선택한 조건에 맞는 문제가 없습니다.")
        return

    # 문제 추가
    for idx, problem in enumerate(problems, start=1):
        doc.add_paragraph(f"문제 {idx}: {problem['problem_text']}")
        if problem['image_path']:
            doc.add_picture(problem['image_path'], width=docx.shared.Inches(4))

    # 워드 파일 저장
    doc.save(output_file)
    print(f"시험지가 생성되었습니다: {output_file}")

if __name__ == "__main__":
    generate_test(chapter="Chapter 1", difficulty="easy", output_file="output/new_test.docx")
